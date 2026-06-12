#!/usr/bin/env python3
"""Format X:\\A\\111.docx per Hangul (한글) Oral Presentation style sheet → 333.docx (+ 333.hwp)."""

from __future__ import annotations

import hashlib
import re
import sys
import zipfile
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from xml.etree import ElementTree as ET

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

INPUT_PATH = Path(r"X:\A\111.docx")
OUTPUT_DOCX = Path(r"X:\A\333.docx")
OUTPUT_HWP = Path(r"X:\A\333.hwp")
AUDIT_PATH = Path(r"X:\A\333_audit.md")

MARGINS = dict(top=Cm(2.5), bottom=Cm(2.5), left=Cm(2.5), right=Cm(2.5))

# Hangul display style names (styles.xml w:name)
HANGUL_BODY = "본문"
HANGUL_TITLE = "논문제목"
HANGUL_SUBTITLE = "논문부제목"
HANGUL_EN_TITLE = "논문영어제목"
HANGUL_AUTHOR = "이름"
HANGUL_AFFILIATION = "소속 및 직책"
HANGUL_H1 = "본문제목1."
HANGUL_H2 = "본문소제목1.1."
HANGUL_H3 = "본문소제목1.1.1."
HANGUL_REFS_TITLE = "참고문헌제목"
HANGUL_REFS = "참고문헌"
HANGUL_FIG_CAP = "그림캡션"
HANGUL_TBL_CAP = "표캡션"
HANGUL_TBL_HEAD = "표제목"
HANGUL_TBL_BODY = "표내용"
HANGUL_FOOTNOTE = "각주"

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

REVERSE_ALIGN = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    None: "inherit",
}


@dataclass
class FormatSpec:
    font: str
    size: float | None = None
    scale: int | None = None
    csp: int | None = None
    line: float | None = None
    align: str = "left"
    bold: bool | None = None
    style_name: str | None = None


SPEC: dict[str, FormatSpec] = {
    "Paper Title": FormatSpec(
        font="Nanum Gothic ExtraBold",
        size=18,
        scale=90,
        csp=-10,
        line=1.3,
        align="center",
        bold=True,
        style_name=HANGUL_TITLE,
    ),
    "Paper Subtitle": FormatSpec(
        font="Nanum Gothic",
        size=11,
        scale=100,
        csp=-10,
        line=1.3,
        align="center",
        style_name=HANGUL_SUBTITLE,
    ),
    "Paper English Title": FormatSpec(
        font="Nanum Gothic ExtraBold",
        size=11,
        scale=100,
        csp=-10,
        line=1.3,
        align="center",
        bold=True,
        style_name=HANGUL_EN_TITLE,
    ),
    "Paper Author": FormatSpec(
        font="Nanum Gothic",
        size=12,
        scale=95,
        csp=-5,
        line=1.6,
        align="center",
        style_name=HANGUL_AUTHOR,
    ),
    "Paper Affiliation": FormatSpec(
        font="Nanum Gothic",
        size=9.5,
        scale=95,
        csp=-5,
        line=1.6,
        align="center",
        style_name=HANGUL_AFFILIATION,
    ),
    "body_h1": FormatSpec(
        font="Nanum Myeongjo",
        size=13,
        scale=100,
        csp=-5,
        line=1.6,
        align="left",
        style_name=HANGUL_H1,
    ),
    "body_h2": FormatSpec(
        font="Nanum Gothic",
        size=11,
        scale=100,
        csp=-5,
        line=1.6,
        align="left",
        style_name=HANGUL_H2,
    ),
    "body_h3": FormatSpec(
        font="Nanum Gothic Bold",
        size=10,
        scale=100,
        csp=-5,
        line=1.6,
        align="left",
        bold=True,
        style_name=HANGUL_H3,
    ),
    "body": FormatSpec(
        font="Nanum Myeongjo",
        size=10,
        scale=100,
        csp=-5,
        line=1.6,
        align="justify",
        style_name=HANGUL_BODY,
    ),
    "refs_title": FormatSpec(
        font="Nanum Gothic ExtraBold",
        size=11,
        scale=100,
        csp=-5,
        line=1.6,
        align="left",
        bold=True,
        style_name=HANGUL_REFS_TITLE,
    ),
    "refs": FormatSpec(
        font="Nanum Myeongjo",
        size=9,
        scale=100,
        csp=None,
        line=1.6,
        align="left",
        style_name=HANGUL_REFS,
    ),
    "figure_caption": FormatSpec(
        font="Nanum Gothic",
        size=8,
        scale=95,
        csp=-5,
        line=1.3,
        align="center",
        style_name=HANGUL_FIG_CAP,
    ),
    "table_caption": FormatSpec(
        font="Nanum Gothic",
        size=8,
        scale=95,
        csp=-5,
        line=1.3,
        align="left",
        style_name=HANGUL_TBL_CAP,
    ),
    "table_header": FormatSpec(
        font="Nanum Myeongjo",
        size=8.5,
        scale=100,
        csp=-5,
        line=1.3,
        align="center",
        bold=True,
        style_name=HANGUL_TBL_HEAD,
    ),
    "table_body": FormatSpec(
        font="Nanum Myeongjo",
        size=8,
        scale=95,
        csp=-10,
        line=1.3,
        align="left",
        style_name=HANGUL_TBL_BODY,
    ),
    "footnote": FormatSpec(
        font="Nanum Myeongjo",
        size=9,
        scale=95,
        csp=-10,
        line=1.6,
        align="justify",
        style_name=HANGUL_FOOTNOTE,
    ),
}

# Legacy style names in 111.docx → spec key
LEGACY_STYLE_MAP = {
    "Paper Title": "Paper Title",
    "Paper Subtitle": "Paper Subtitle",
    "Paper Author": "Paper Author",
    "Paper Affiliation": "Paper Affiliation",
    "Paper H1": "body_h1",
    "Paper H2": "body_h2",
    "Paper H3": "body_h3",
    "Paper Ref": "refs",
    "样式1": "body_h2",
    "样式2": "body_h1",
    "样式3": "body",
    "Caption": "figure_caption",
    "List Paragraph": "body",
    "Normal": "body",
}

# styles.xml: styleId or display name → FormatSpec
STYLE_XML_DEFS: dict[str, FormatSpec] = {
    "PaperTitle": SPEC["Paper Title"],
    "PaperSubtitle": SPEC["Paper Subtitle"],
    "PaperAuthor": SPEC["Paper Author"],
    "PaperAffiliation": SPEC["Paper Affiliation"],
    "PaperH1": SPEC["body_h1"],
    "PaperH2": SPEC["body_h2"],
    "PaperH3": SPEC["body_h3"],
    "PaperRef": SPEC["refs"],
    "1": SPEC["body_h2"],
    "2": SPEC["body_h1"],
    "3": SPEC["body"],
}

STYLE_NAME_DEFS: dict[str, FormatSpec] = {
    **{spec.style_name: spec for spec in SPEC.values() if spec.style_name},
    "Paper Title": SPEC["Paper Title"],
    "Paper Subtitle": SPEC["Paper Subtitle"],
    "Paper Author": SPEC["Paper Author"],
    "Paper Affiliation": SPEC["Paper Affiliation"],
    "Paper H1": SPEC["body_h1"],
    "Paper H2": SPEC["body_h2"],
    "Paper H3": SPEC["body_h3"],
    "Paper Ref": SPEC["refs"],
    "样式1": SPEC["body_h2"],
    "样式2": SPEC["body_h1"],
    "样式3": SPEC["body"],
    "Normal": SPEC["body"],
    "Caption": SPEC["figure_caption"],
}

# styleId → Hangul display name for styles.xml rename
STYLE_ID_HANGUL_NAME: dict[str, str] = {
    "PaperTitle": HANGUL_TITLE,
    "PaperSubtitle": HANGUL_SUBTITLE,
    "PaperAuthor": HANGUL_AUTHOR,
    "PaperAffiliation": HANGUL_AFFILIATION,
    "PaperH1": HANGUL_H1,
    "PaperH2": HANGUL_H2,
    "PaperH3": HANGUL_H3,
    "PaperRef": HANGUL_REFS,
    "Normal": HANGUL_BODY,
    "Caption": HANGUL_FIG_CAP,
}


def wqn(tag: str) -> str:
    return f"{{{WNS}}}{tag}"


def paragraph_text_hash(doc: Document) -> str:
    return hashlib.sha256("\n".join(p.text for p in doc.paragraphs).encode("utf-8")).hexdigest()


def is_image_only_paragraph(paragraph) -> bool:
    el = paragraph._element
    has_drawing = el.find(".//" + wqn("drawing")) is not None
    has_pict = el.find(".//" + wqn("pict")) is not None
    return (has_drawing or has_pict) and not paragraph.text.strip()


def is_figure_caption_pattern(text: str) -> bool:
    t = text.strip()
    return bool(
        re.match(r"^\[?\s*그림\s*\d+\s*\]?", t, re.IGNORECASE)
        or re.match(r"^\[그림\s*\d+\]", t)
        or re.match(r"^<\s*그림\s*\d+\s*>", t)
        or re.match(r"^图\s*\d+", t)
    )


def is_table_caption_pattern(text: str) -> bool:
    t = text.strip()
    return bool(re.match(r"^<\s*표\s*\d+\s*>", t) or re.match(r"^표\s*\d+", t) or re.match(r"^表\s*\d+", t))


def is_adjacent_to_image_only(doc: Document, index: int) -> bool:
    for delta in (-1, 1):
        j = index + delta
        if 0 <= j < len(doc.paragraphs) and is_image_only_paragraph(doc.paragraphs[j]):
            return True
    return False


def is_likely_figure_caption(doc: Document, index: int, text: str) -> bool:
    if is_figure_caption_pattern(text):
        return True
    if heading_level(text) is not None:
        return False
    if not is_adjacent_to_image_only(doc, index):
        return False
    return len(text) <= 120


def heading_level(text: str) -> int | None:
    m = re.match(r"^(\d+(?:\.\d+)*)\.\s", text.strip())
    if not m:
        return None
    parts = m.group(1).split(".")
    depth = len(parts)
    if depth >= 3:
        return 3
    if depth == 2:
        return 2
    if depth == 1:
        rest = text.strip()[m.end() :]
        if ":" in rest:
            return None
        return 1
    return None


def is_enumerated_body_item(text: str) -> bool:
    if re.match(r"^\d+\.\d+", text.strip()):
        return False
    m = re.match(r"^(\d+)\.\s+", text.strip())
    if not m:
        return False
    rest = text.strip()[m.end() :]
    return ":" in rest and len(text.strip()) > 35


def is_unnumbered_subheading(text: str) -> bool:
    t = text.strip()
    if not t or re.match(r"^\d", t):
        return False
    return len(t) <= 30


def is_author_line(text: str) -> bool:
    t = text.strip()
    return bool(
        re.search(r"LIANG|SHICHENG", t, re.I)
        or re.match(r"^[가-힣A-Za-z\s]+\([A-Za-z]", t)
        or (")" in t and "(" in t and len(t) < 80)
    )


def is_affiliation_line(text: str) -> bool:
    t = text.strip()
    return bool(re.search(r"대학|대학원|University|College|학과|과정", t, re.I))


def classify_paragraph(paragraph, index: int, refs_title_index: int | None, doc: Document) -> str | None:
    if is_image_only_paragraph(paragraph):
        return None

    text = paragraph.text.strip()
    if not text:
        return None

    style = paragraph.style.name if paragraph.style else ""

    if is_table_caption_pattern(text):
        return "table_caption"
    if is_likely_figure_caption(doc, index, text):
        return "figure_caption"

    # Title page block (111.docx uses Paper Subtitle style for title/author lines)
    if index == 0:
        return "Paper Title"
    if index == 1 and is_author_line(text):
        return "Paper Author"
    if index == 2 and not heading_level(text) and not is_affiliation_line(text):
        if style in (HANGUL_H3, HANGUL_SUBTITLE, "Paper Subtitle") or text.startswith("—") or text.startswith("-"):
            return "Paper Subtitle"
    if index <= 4 and is_affiliation_line(text):
        return "Paper Affiliation"

    if style == "Paper Title":
        return "Paper Title"
    if style == "Paper Subtitle" and index <= 2:
        return "Paper Author" if is_author_line(text) else "Paper Subtitle"
    if style == "Paper Author":
        return "Paper Author"
    if style == "Paper Affiliation":
        return "Paper Affiliation"

    if "참고문헌" in text and len(text) < 20:
        return "refs_title"

    if style in (HANGUL_REFS, "Paper Ref"):
        return "refs"

    if refs_title_index is not None and index > refs_title_index:
        if heading_level(text) is None and not is_table_caption_pattern(text):
            return "refs"
        if re.match(r"^\d+\.\s+\S", text):
            return "refs"

    if is_enumerated_body_item(text):
        return "body"
    if is_unnumbered_subheading(text) and refs_title_index is not None and index < refs_title_index:
        return "body_h3"

    level = heading_level(text)
    if level == 1:
        return "body_h1"
    if level == 2:
        return "body_h2"
    if level == 3:
        return "body_h3"

    if style in ("Paper H1", "样式2"):
        return "body_h1"
    if style in ("Paper H2", "样式1"):
        return "body_h2"
    if style in ("Paper H3",):
        return "body_h3"
    if style in ("Paper Ref",):
        return "refs"

    return "body"


def find_refs_title_index(doc: Document) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        if "참고문헌" in p.text.strip() and len(p.text.strip()) < 20:
            return i
    return None


def ensure_rpr(element) -> Any:
    if hasattr(element, "get_or_add_rPr"):
        return element.get_or_add_rPr()
    rpr = element.find(wqn("rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        element.insert(0, rpr)
    return rpr


def set_xml_font_rpr(rpr, font_name: str) -> None:
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for tag in ("ascii", "hAnsi", "eastAsia", "cs"):
        rf.set(qn(f"w:{tag}"), font_name)


def set_xml_size_rpr(rpr, size_pt: float) -> None:
    half = str(int(round(size_pt * 2)))
    for tag in ("sz", "szCs"):
        el = rpr.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            rpr.append(el)
        el.set(qn("w:val"), half)


def set_xml_spacing_rpr(rpr, csp: int | None) -> None:
    sp = rpr.find(qn("w:spacing"))
    if csp is None:
        if sp is not None:
            rpr.remove(sp)
        return
    if sp is None:
        sp = OxmlElement("w:spacing")
        rpr.append(sp)
    sp.set(qn("w:val"), str(csp))


def set_xml_scale_rpr(rpr, scale: int | None) -> None:
    for tag in ("textScale",):
        el = rpr.find(qn(f"w:{tag}"))
        if el is not None:
            rpr.remove(el)
    w_el = rpr.find(qn("w:w"))
    if scale is None:
        if w_el is not None:
            rpr.remove(w_el)
        return
    if w_el is None:
        w_el = OxmlElement("w:w")
        rpr.append(w_el)
    w_el.set(qn("w:val"), str(scale))


def set_xml_bold_rpr(rpr, bold: bool | None) -> None:
    b = rpr.find(qn("w:b"))
    if bold is None:
        if b is not None:
            rpr.remove(b)
        return
    if b is None:
        b = OxmlElement("w:b")
        rpr.append(b)
    if not bold:
        b.set(qn("w:val"), "0")


JC_XML = {"left": "left", "center": "center", "right": "right", "justify": "both"}


def ensure_ppr(paragraph) -> Any:
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        paragraph._element.insert(0, p_pr)
    return p_pr


def apply_paragraph_ppr_xml(paragraph, spec: FormatSpec) -> None:
    p_pr = ensure_ppr(paragraph)
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    jc.set(qn("w:val"), JC_XML[spec.align])

    if spec.line is not None:
        sp = p_pr.find(qn("w:spacing"))
        if sp is None:
            sp = OxmlElement("w:spacing")
            p_pr.append(sp)
        sp.set(qn("w:line"), str(int(round(spec.line * 240))))
        sp.set(qn("w:lineRule"), "auto")


def apply_run_format(run, spec: FormatSpec) -> None:
    rpr = ensure_rpr(run._element)
    set_xml_font_rpr(rpr, spec.font)
    if spec.size is not None:
        set_xml_size_rpr(rpr, spec.size)
        run.font.size = Pt(spec.size)
    run.font.name = spec.font
    set_xml_spacing_rpr(rpr, spec.csp)
    set_xml_scale_rpr(rpr, spec.scale)
    if spec.bold is not None:
        set_xml_bold_rpr(rpr, spec.bold)
        run.font.bold = spec.bold


def apply_paragraph_format(paragraph, spec: FormatSpec) -> None:
    pf = paragraph.paragraph_format
    pf.alignment = ALIGN_MAP[spec.align]
    if spec.line is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = spec.line
    apply_paragraph_ppr_xml(paragraph, spec)

    if spec.style_name:
        try:
            paragraph.style = spec.style_name
        except KeyError:
            for legacy in ("Paper H1", "Paper H2", "Paper H3", "Normal"):
                try:
                    paragraph.style = legacy
                    break
                except KeyError:
                    continue

    from docx.text.run import Run

    for r_el in paragraph._element.findall(".//" + wqn("r")):
        apply_run_format(Run(r_el, paragraph), spec)


def apply_cell_format(cell, spec: FormatSpec) -> None:
    for p in cell.paragraphs:
        apply_paragraph_format(p, spec)


def format_tables(doc: Document) -> list[str]:
    changes: list[str] = []
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            spec = SPEC["table_header"] if ri == 0 else SPEC["table_body"]
            for cell in row.cells:
                apply_cell_format(cell, spec)
            changes.append(f"Table {ti} row {ri}: {spec.style_name}")
    return changes


def ensure_hangul_styles(doc: Document) -> None:
    """Create paragraph styles with Hangul display names if missing."""
    existing = {s.name for s in doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH}
    base = doc.styles["Normal"]
    for key, spec in SPEC.items():
        if not spec.style_name or spec.style_name in existing:
            continue
        try:
            new_style = doc.styles.add_style(spec.style_name, WD_STYLE_TYPE.PARAGRAPH)
            new_style.base_style = base
            existing.add(spec.style_name)
        except ValueError:
            pass


def setup_margins(doc: Document) -> None:
    for sec in doc.sections:
        sec.top_margin = MARGINS["top"]
        sec.bottom_margin = MARGINS["bottom"]
        sec.left_margin = MARGINS["left"]
        sec.right_margin = MARGINS["right"]


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    r = run._r
    for ftype in ("begin", "separate", "end"):
        if ftype == "begin":
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            r.append(fld_begin)
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = " PAGE "
            r.append(instr)
            fld_sep = OxmlElement("w:fldChar")
            fld_sep.set(qn("w:fldCharType"), "separate")
            r.append(fld_sep)
            t = OxmlElement("w:t")
            t.text = "1"
            r.append(t)
        else:
            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            r.append(fld_end)
    run.font.name = "Nanum Gothic"
    run.font.size = Pt(9)


def setup_page_footers(doc: Document) -> None:
    for sec in doc.sections:
        footer = sec.footer
        footer.is_linked_to_previous = False
        root = footer._element
        for child in list(root):
            if child.tag == wqn("p"):
                root.remove(child)
        p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_prefix = p.add_run("- ")
        r_prefix.font.name = "Nanum Gothic"
        r_prefix.font.size = Pt(9)
        add_page_number_field(p)
        r_suffix = p.add_run(" -")
        r_suffix.font.name = "Nanum Gothic"
        r_suffix.font.size = Pt(9)


def patch_first_section_footer_reference(docx_path: Path) -> None:
    with zipfile.ZipFile(docx_path, "r") as zin:
        contents = {name: zin.read(name) for name in zin.namelist()}

    root = ET.fromstring(contents["word/document.xml"])
    body = root.find(wqn("body"))
    if body is None:
        return

    last_sp = body.find(wqn("sectPr"))
    if last_sp is None:
        return
    footer_ref = last_sp.find(wqn("footerReference"))
    if footer_ref is None:
        return

    for p in body.findall(wqn("p")):
        p_pr = p.find(wqn("pPr"))
        if p_pr is None:
            continue
        sp = p_pr.find(wqn("sectPr"))
        if sp is None:
            continue
        if sp.find(wqn("footerReference")) is None:
            sp.append(deepcopy(footer_ref))
        break

    contents["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    tmp = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in contents.items():
            zout.writestr(name, data)
    tmp.replace(docx_path)


def build_style_rpr(spec: FormatSpec) -> ET.Element:
    rpr = ET.Element(wqn("rPr"))
    rf = ET.SubElement(rpr, wqn("rFonts"))
    rf.set(wqn("ascii"), spec.font)
    rf.set(wqn("hAnsi"), spec.font)
    rf.set(wqn("eastAsia"), spec.font)
    if spec.size is not None:
        half = str(int(round(spec.size * 2)))
        sz = ET.SubElement(rpr, wqn("sz"))
        sz.set(wqn("val"), half)
        szcs = ET.SubElement(rpr, wqn("szCs"))
        szcs.set(wqn("val"), half)
    if spec.csp is not None:
        sp = ET.SubElement(rpr, wqn("spacing"))
        sp.set(wqn("val"), str(spec.csp))
    if spec.scale is not None:
        w_el = ET.SubElement(rpr, wqn("w"))
        w_el.set(wqn("val"), str(spec.scale))
    if spec.bold:
        ET.SubElement(rpr, wqn("b"))
    return rpr


def build_style_ppr(spec: FormatSpec) -> ET.Element:
    ppr = ET.Element(wqn("pPr"))
    jc = ET.SubElement(ppr, wqn("jc"))
    jc.set(wqn("val"), JC_XML[spec.align])
    if spec.line is not None:
        sp = ET.SubElement(ppr, wqn("spacing"))
        sp.set(wqn("line"), str(int(spec.line * 240)))
        sp.set(wqn("lineRule"), "auto")
    return ppr


def update_styles_xml(styles_xml_bytes: bytes) -> bytes:
    root = ET.fromstring(styles_xml_bytes)

    for st in root.findall(wqn("style")):
        name_el = st.find(wqn("name"))
        if name_el is None:
            continue
        name = name_el.get(wqn("val"), "")
        sid = st.get(wqn("styleId"), "")

        spec = STYLE_NAME_DEFS.get(name) or STYLE_XML_DEFS.get(sid)
        if spec is None:
            continue

        hangul_name = STYLE_ID_HANGUL_NAME.get(sid) or spec.style_name
        if hangul_name:
            name_el.set(wqn("val"), hangul_name)

        for tag in ("rPr", "pPr"):
            old = st.find(wqn(tag))
            if old is not None:
                st.remove(old)

        st.insert(0, build_style_ppr(spec))
        st.insert(1, build_style_rpr(spec))

    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def get_run_props(run) -> dict[str, Any]:
    rpr = run._element.rPr
    d: dict[str, Any] = {}
    if run.font.name:
        d["font"] = run.font.name
    if run.font.size:
        d["size"] = round(run.font.size.pt, 1)
    if run.font.bold:
        d["bold"] = True
    if rpr is not None:
        sp = rpr.find(qn("w:spacing"))
        if sp is not None and sp.get(qn("w:val")) is not None:
            d["csp"] = int(sp.get(qn("w:val")))
        w = rpr.find(qn("w:w"))
        if w is not None and w.get(qn("w:val")) is not None:
            d["scale"] = int(w.get(qn("w:val")))
        ts = rpr.find(qn("w:textScale"))
        if ts is not None:
            d["textScale"] = ts.get(qn("w:val"))
    return d


def get_para_props(paragraph) -> dict[str, Any]:
    pf = paragraph.paragraph_format
    props: dict[str, Any] = {"align": REVERSE_ALIGN.get(pf.alignment), "line": pf.line_spacing}
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is not None:
        jc = p_pr.find(qn("w:jc"))
        if jc is not None:
            props["align"] = {"left": "left", "center": "center", "right": "right", "both": "justify"}.get(
                jc.get(qn("w:val")), jc.get(qn("w:val"))
            )
        sp = p_pr.find(qn("w:spacing"))
        if sp is not None and sp.get(qn("w:line")):
            props["line"] = round(int(sp.get(qn("w:line"))) / 240, 2)
    return props


def audit_document(doc: Document) -> list[tuple]:
    issues = []
    refs_idx = find_refs_title_index(doc)
    for i, p in enumerate(doc.paragraphs):
        if is_image_only_paragraph(p) or not p.text.strip():
            continue
        cat = classify_paragraph(p, i, refs_idx, doc)
        if cat is None or cat not in SPEC:
            continue
        spec = SPEC[cat]
        rp = get_run_props(p.runs[0]) if p.runs else {}
        pp = get_para_props(p)

        if spec.bold is not None and rp.get("bold") != spec.bold:
            issues.append((i, cat, "bold", spec.bold, rp.get("bold")))

        for key, expected in [
            ("font", spec.font.split()[0]),
            ("size", spec.size),
            ("scale", spec.scale if spec.scale != 100 else 100),
            ("csp", spec.csp),
            ("line", spec.line),
            ("align", spec.align),
        ]:
            if key == "font":
                actual = rp.get("font", "")
                if expected and actual and expected not in actual:
                    issues.append((i, cat, key, expected, actual))
            elif key == "csp" and expected is None:
                continue
            elif key == "scale" and expected == 100:
                if rp.get("scale") not in (None, 100) or rp.get("textScale"):
                    issues.append((i, cat, key, 100, rp.get("scale") or rp.get("textScale")))
            elif key == "scale" and expected != 100:
                if rp.get("scale") != expected:
                    issues.append((i, cat, key, expected, rp.get("scale", "missing")))
            elif key in rp:
                if rp[key] != expected:
                    issues.append((i, cat, key, expected, rp[key]))
            elif key in pp:
                actual = pp[key]
                if key == "align" and (actual in ("inherit", None) or actual != expected):
                    issues.append((i, cat, key, expected, actual))
                elif key == "line" and (actual is None or abs(float(actual) - float(expected)) > 0.05):
                    issues.append((i, cat, key, expected, actual))
            elif key in ("size", "scale", "csp") and expected is not None:
                issues.append((i, cat, key, expected, "missing"))
            elif key == "line" and pp.get("line") is None:
                issues.append((i, cat, key, expected, None))
            elif key == "align" and pp.get("align") in ("inherit", None):
                issues.append((i, cat, key, expected, pp.get("align")))
    return issues


def format_document(doc: Document) -> list[str]:
    changes: list[str] = []
    ensure_hangul_styles(doc)
    setup_margins(doc)

    refs_idx = find_refs_title_index(doc)
    for i, p in enumerate(doc.paragraphs):
        if is_image_only_paragraph(p) or not p.text.strip():
            continue
        cat = classify_paragraph(p, i, refs_idx, doc)
        if cat is None:
            continue
        spec = SPEC[cat]
        old_style = p.style.name if p.style else ""
        apply_paragraph_format(p, spec)
        new_style = p.style.name if p.style else ""
        changes.append(
            f"Para {i}: {old_style!r} -> {new_style!r} ({cat})"
            if old_style != new_style
            else f"Para {i}: applied {cat}"
        )

    changes.extend(format_tables(doc))
    setup_page_footers(doc)
    return changes


def media_hash(docx_path: Path) -> dict[str, str]:
    out: dict[str, str] = {}
    with zipfile.ZipFile(docx_path, "r") as z:
        for name in z.namelist():
            if name.startswith("word/media/"):
                out[name] = hashlib.sha256(z.read(name)).hexdigest()
    return out


def check_fonts() -> list[str]:
    missing: list[str] = []
    required = [
        "Nanum Gothic",
        "Nanum Gothic ExtraBold",
        "Nanum Gothic Bold",
        "Nanum Myeongjo",
    ]
    try:
        import winreg

        key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts")
        installed = " ".join(winreg.EnumValue(key, i)[0] for i in range(winreg.QueryInfoKey(key)[1])).lower()
        for font in required:
            if font.lower().replace(" ", "") not in installed.replace(" ", ""):
                missing.append(font)
    except OSError:
        pass
    return missing


def find_hancom() -> Path | None:
    candidates = [
        Path(r"C:\Program Files\Hancom\HancomOffice\HOffice130\Bin\Hwp.exe"),
        Path(r"C:\Program Files (x86)\Hancom\HancomOffice\HOffice130\Bin\Hwp.exe"),
        Path(r"C:\Program Files\Hancom\HancomOffice\HOffice120\Bin\Hwp.exe"),
        Path(r"C:\Program Files (x86)\Hancom\HancomOffice\HOffice120\Bin\Hwp.exe"),
    ]
    for p in candidates:
        if p.exists():
            return p
    return None


def export_hwp_via_com(docx_path: Path, hwp_path: Path) -> tuple[bool, str]:
    try:
        import win32com.client  # type: ignore
    except ImportError:
        return False, "pywin32 未安装，无法调用 Hancom COM"

    prog_ids = ("HWPFrame.HwpObject", "HWPFrame.HwpObject.1", "HWPApplication.HwpObject")
    last_err = ""
    for prog_id in prog_ids:
        try:
            hwp = win32com.client.Dispatch(prog_id)
            try:
                hwp.RegisterModule("FilePathCheckDLL", "SecurityModule")
            except Exception:
                pass
            hwp.Open(str(docx_path))
            hwp.SaveAs(str(hwp_path), "HWP")
            hwp.Quit()
            if hwp_path.exists():
                return True, f"已通过 {prog_id} 保存 {hwp_path}"
            last_err = f"{prog_id}: SaveAs 完成但文件未找到"
        except Exception as exc:
            last_err = f"{prog_id}: {exc}"
            continue
    return False, last_err or "Hancom Office COM 不可用"


def write_audit_report(
    before_issues: list,
    after_issues: list,
    changes: list[str],
    text_hash_before: str,
    text_hash_after: str,
    missing_fonts: list[str],
    hwp_result: tuple[bool, str],
) -> None:
    hwp_ok, hwp_msg = hwp_result
    lines = [
        "# 333.docx 한글样式格式化报告",
        "",
        "## 摘要",
        "",
        f"- 输入: `{INPUT_PATH}`",
        f"- 输出: `{OUTPUT_DOCX}`",
        f"- HWP: `{OUTPUT_HWP}` — {'成功' if hwp_ok else '未生成'}",
        f"- 修正前问题数: **{len(before_issues)}**",
        f"- 修正后问题数: **{len(after_issues)}**",
        f"- 文本 hash 一致: **{'是' if text_hash_before == text_hash_after else '否'}**",
        "",
        "## 장평 / 자간 / 행간 / 정렬 核对标准",
        "",
        "| 样式 | 장평 (w:w) | 자간 | 행간 | 정렬 |",
        "|------|------------|------|------|------|",
        "| 본문 | 100% | -5 | 160% | 양쪽 |",
        "| 논문제목 | 90% | -10 | 130% | 중앙 |",
        "| 논문부제목 | 100% | -10 | 130% | 중앙 |",
        "| 이름/소속 | 95% | -5 | 160% | 중앙 |",
        "| 본문제목1. | 100% | -5 | 160% | 좌측 |",
        "| 본문소제목1.1. | 100% | -5 | 160% | 좌측 |",
        "| 본문소제목1.1.1. | 100% | -5 | 160% | 좌측 (Bold) |",
        "| 그림캡션 | 95% | -5 | 130% | 중앙 |",
        "| 표캡션 | 95% | -5 | 130% | 좌측 |",
        "| 표제목 | 100% | -5 | 130% | 중앙 (진하게) |",
        "| 표내용 | 95% | -10 | 130% | 좌측 |",
        "| 참고문헌제목 | 100% | -5 | 160% | 좌측 (고딕 ExtraBold) |",
        "| 참고문헌 | 100% | (空) | 160% | 좌측 |",
        "",
        "> 注：图3 标注 참고문헌제목为 나눔명조 ExtraBold，官方样式表为 나눔고딕 ExtraBold，本脚本按样式表执行。",
        "",
        "## 字体检查",
        "",
    ]
    if missing_fonts:
        lines.append("以下字体可能未安装，请在 [네이버 글꼴](https://hangeul.naver.com/font) 下载：")
        for f in missing_fonts:
            lines.append(f"- {f}")
    else:
        lines.append("- Nanum Gothic / Myeongjo 系列字体已检测到")

    lines.extend(["", "## HWP 导出", "", f"- {hwp_msg}"])
    if not hwp_ok:
        lines.extend(
            [
                "",
                "### 手动转换步骤",
                "1. 用 **한글(Hangul)** 打开 `X:\\A\\333.docx`",
                "2. 「파일」→「다른 이름으로 저장」→ 格式选 **HWP** → 保存为 `333.hwp`",
                "3. 在 한글 中目视核对 장평/자간/행간 是否与 Oral Presentation 模板一致",
            ]
        )

    lines.extend(["", "## 格式化操作", ""])
    for c in changes[:60]:
        lines.append(f"- {c}")
    if len(changes) > 60:
        lines.append(f"- ... 共 {len(changes)} 项")

    lines.extend(["", "## 修正后剩余问题", ""])
    if after_issues:
        for issue in after_issues[:50]:
            lines.append(f"- Para {issue[0]} [{issue[1]}] {issue[2]}: expected {issue[3]}, got {issue[4]}")
        if len(after_issues) > 50:
            lines.append(f"- ... 共 {len(after_issues)} 项")
    else:
        lines.append("- 无（段落级全部合规）")

    AUDIT_PATH.write_text("\n".join(lines), encoding="utf-8")


def main() -> int:
    if not INPUT_PATH.exists():
        print(f"Input not found: {INPUT_PATH}", file=sys.stderr)
        return 1

    missing_fonts = check_fonts()
    media_before = media_hash(INPUT_PATH)

    doc = Document(str(INPUT_PATH))
    hash_before = paragraph_text_hash(doc)
    before_issues = audit_document(doc)

    changes = format_document(doc)
    hash_after = paragraph_text_hash(doc)

    if hash_before != hash_after:
        print("ERROR: Text content changed!", file=sys.stderr)
        return 2

    doc.save(str(OUTPUT_DOCX))

    with zipfile.ZipFile(OUTPUT_DOCX, "r") as zin:
        contents = {name: zin.read(name) for name in zin.namelist()}
    if "word/styles.xml" in contents:
        contents["word/styles.xml"] = update_styles_xml(contents["word/styles.xml"])
    tmp = OUTPUT_DOCX.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in contents.items():
            zout.writestr(name, data)
    tmp.replace(OUTPUT_DOCX)

    patch_first_section_footer_reference(OUTPUT_DOCX)

    media_after = media_hash(OUTPUT_DOCX)
    if media_before != media_after:
        print("ERROR: Media files changed!", file=sys.stderr)
        return 4

    doc_out = Document(str(OUTPUT_DOCX))
    after_issues = audit_document(doc_out)

    hwp_result = (False, "Hancom Office 未检测到")
    if find_hancom():
        hwp_result = export_hwp_via_com(OUTPUT_DOCX, OUTPUT_HWP)
    else:
        try:
            import win32com.client  # type: ignore

            hwp_result = export_hwp_via_com(OUTPUT_DOCX, OUTPUT_HWP)
        except ImportError:
            pass

    write_audit_report(before_issues, after_issues, changes, hash_before, hash_after, missing_fonts, hwp_result)

    print(f"Saved: {OUTPUT_DOCX}")
    print(f"Audit: {AUDIT_PATH}")
    print(f"Before issues: {len(before_issues)}, After issues: {len(after_issues)}")
    print(f"HWP: {hwp_result[1]}")
    return 0 if len(after_issues) == 0 else 3


if __name__ == "__main__":
    sys.exit(main())
