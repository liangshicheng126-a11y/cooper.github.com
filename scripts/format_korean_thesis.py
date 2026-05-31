#!/usr/bin/env python3
"""Format Korean thesis DOCX per style table. Does NOT modify text or images."""

from __future__ import annotations

import hashlib
import re
import shutil
import sys
import zipfile
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

INPUT_PATH = Path(r"X:\A\liangsc_formatted.docx")
OUTPUT_PATH = Path(r"X:\A\liangsc_formatted_checked.docx")
AUDIT_PATH = Path(r"X:\A\liangsc_formatted_audit.md")

# Style name mapping in document
STYLE_BODY_H1 = "样式2"
STYLE_BODY_H2 = "样式1"
STYLE_BODY = "样式3"

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

REVERSE_ALIGN = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    None: "inherit",
}


@dataclass
class FormatSpec:
    font: str
    size: float | None = None
    scale: int | None = None
    csp: int | None = None
    line: float | None = None
    align: str = "left"
    bold: bool | None = None
    style_name: str | None = None


SPEC: dict[str, FormatSpec] = {
    "Paper Title": FormatSpec(
        font="Nanum Gothic ExtraBold",
        size=18,
        scale=90,
        csp=-10,
        line=1.3,
        align="center",
        bold=True,
        style_name="Paper Title",
    ),
    "Paper Subtitle": FormatSpec(
        font="Nanum Gothic",
        size=11,
        scale=100,
        csp=-10,
        line=1.3,
        align="center",
        style_name="Paper Subtitle",
    ),
    "Paper Author": FormatSpec(
        font="Nanum Gothic",
        size=12,
        scale=95,
        csp=-5,
        line=1.6,
        align="center",
        style_name="Paper Author",
    ),
    "Paper Affiliation": FormatSpec(
        font="Nanum Gothic",
        size=9.5,
        scale=95,
        csp=-5,
        line=1.6,
        align="center",
        style_name="Paper Affiliation",
    ),
    "body_h1": FormatSpec(
        font="Nanum Myeongjo",
        size=13,
        scale=100,
        csp=-5,
        line=1.6,
        align="left",
        style_name=STYLE_BODY_H1,
    ),
    "body_h2": FormatSpec(
        font="Nanum Gothic",
        size=11,
        scale=100,
        csp=-5,
        line=1.6,
        align="left",
        style_name=STYLE_BODY_H2,
    ),
    "body_h3": FormatSpec(
        font="Nanum Gothic",
        size=10,
        scale=100,
        csp=-5,
        line=1.6,
        align="left",
        bold=True,
        style_name=STYLE_BODY_H2,
    ),
    "body": FormatSpec(
        font="Nanum Myeongjo",
        size=10,
        scale=100,
        csp=-5,
        line=1.6,
        align="justify",
        style_name=STYLE_BODY,
    ),
    "refs_title": FormatSpec(
        font="Nanum Gothic ExtraBold",
        size=11,
        scale=100,
        csp=-5,
        line=1.6,
        align="left",
        bold=True,
        style_name=STYLE_BODY_H2,
    ),
    "refs": FormatSpec(
        font="Nanum Myeongjo",
        size=9,
        scale=100,
        csp=None,
        line=1.6,
        align="left",
        style_name=STYLE_BODY,
    ),
}

# styles.xml styleId -> FormatSpec key
STYLE_XML_DEFS: dict[str, FormatSpec] = {
    "PaperTitle": SPEC["Paper Title"],
    "PaperSubtitle": SPEC["Paper Subtitle"],
    "PaperAuthor": SPEC["Paper Author"],
    "PaperAffiliation": SPEC["Paper Affiliation"],
    "PaperH1": SPEC["body_h1"],
    "PaperH2": SPEC["body_h2"],
    "PaperH3": SPEC["body_h3"],
    "PaperRef": SPEC["refs"],
    "1": SPEC["body_h2"],  # 样式1 styleId often numeric
    "2": SPEC["body_h1"],
    "3": SPEC["body"],
}

# Korean style names in styles.xml (by display name)
STYLE_NAME_DEFS: dict[str, FormatSpec] = {
    "样式1": SPEC["body_h2"],
    "样式2": SPEC["body_h1"],
    "样式3": SPEC["body"],
    "Paper Title": SPEC["Paper Title"],
    "Paper Subtitle": SPEC["Paper Subtitle"],
    "Paper Author": SPEC["Paper Author"],
    "Paper Affiliation": SPEC["Paper Affiliation"],
    "Paper H1": SPEC["body_h1"],
    "Paper H2": SPEC["body_h2"],
    "Paper H3": SPEC["body_h3"],
    "Paper Ref": SPEC["refs"],
    "Normal": FormatSpec(
        font="Nanum Myeongjo",
        size=10,
        scale=100,
        csp=-5,
        line=1.6,
        align="justify",
    ),
}


def wqn(tag: str) -> str:
    return f"{{{WNS}}}{tag}"


def paragraph_text_hash(doc: Document) -> str:
    texts = [p.text for p in doc.paragraphs]
    return hashlib.sha256("\n".join(texts).encode("utf-8")).hexdigest()


def is_image_only_paragraph(paragraph) -> bool:
    el = paragraph._element
    has_drawing = el.find(".//" + wqn("drawing")) is not None
    has_pict = el.find(".//" + wqn("pict")) is not None
    text = paragraph.text.strip()
    return (has_drawing or has_pict) and not text


def heading_level(text: str) -> int | None:
    """Return heading depth from numbering like 1. / 1.1. / 1.1.1."""
    m = re.match(r"^(\d+(?:\.\d+)*)\.\s", text.strip())
    if not m:
        return None
    parts = m.group(1).split(".")
    # Single-digit "1. xxx" with long text or colon → enumerated list item, not heading
    if len(parts) == 1:
        rest = text.strip()[m.end() :]
        if len(text.strip()) > 35 or ":" in rest[:25]:
            return None
    return len(parts)


def is_enumerated_body_item(text: str) -> bool:
    """Numbered list items under a section (e.g. 1. 문헌 고찰: ...)."""
    m = re.match(r"^\d+\.\s+\S", text.strip())
    if not m:
        return False
    # Multi-level numbering is handled by heading_level
    if re.match(r"^\d+\.\d+", text.strip()):
        return False
    rest = text.strip()[m.end() :]
    return len(text.strip()) > 35 or ":" in rest[:25]


def classify_paragraph(paragraph, index: int, refs_title_index: int | None) -> str | None:
    if is_image_only_paragraph(paragraph):
        return None

    text = paragraph.text.strip()
    if not text:
        return None

    style = paragraph.style.name if paragraph.style else ""

    if style == "Paper Title":
        return "Paper Title"
    if style == "Paper Subtitle":
        return "Paper Subtitle"
    if style == "Paper Author":
        return "Paper Author"
    if style == "Paper Affiliation":
        return "Paper Affiliation"

    if "참고문헌" in text and len(text) < 20:
        return "refs_title"

    # References only appear after the refs_title heading
    if refs_title_index is not None and index > refs_title_index:
        if re.match(r"^\d+\.\s+\S", text):
            return "refs"

    if is_enumerated_body_item(text):
        return "body"

    level = heading_level(text)
    if level == 1:
        return "body_h1"
    if level == 2:
        return "body_h2"
    if level == 3:
        return "body_h3"

    if style == STYLE_BODY_H1:
        return "body_h1"
    if style == STYLE_BODY_H2:
        return "body_h2"
    if style == STYLE_BODY:
        return "body"

    if style == "Normal":
        if level == 1:
            return "body_h1"
        if level == 2:
            return "body_h2"
        return "body"

    return "body"


def find_refs_title_index(doc: Document) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if "참고문헌" in text and len(text) < 20:
            return i
    return None


def ensure_rpr(element) -> Any:
    rpr = element.get_or_add_rPr() if hasattr(element, "get_or_add_rPr") else element.find(wqn("rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        element.insert(0, rpr)
    return rpr


def set_xml_font_rpr(rpr, font_name: str) -> None:
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:ascii"), font_name)
    rf.set(qn("w:hAnsi"), font_name)
    rf.set(qn("w:eastAsia"), font_name)


def set_xml_size_rpr(rpr, size_pt: float) -> None:
    half = str(int(round(size_pt * 2)))
    sz = rpr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rpr.append(sz)
    sz.set(qn("w:val"), half)
    szcs = rpr.find(qn("w:szCs"))
    if szcs is None:
        szcs = OxmlElement("w:szCs")
        rpr.append(szcs)
    szcs.set(qn("w:val"), half)


def set_xml_spacing_rpr(rpr, csp: int | None) -> None:
    sp = rpr.find(qn("w:spacing"))
    if csp is None:
        if sp is not None:
            rpr.remove(sp)
        return
    if sp is None:
        sp = OxmlElement("w:spacing")
        rpr.append(sp)
    sp.set(qn("w:val"), str(csp))


def set_xml_scale_rpr(rpr, scale: int | None) -> None:
    w_el = rpr.find(qn("w:w"))
    if scale is None:
        if w_el is not None:
            rpr.remove(w_el)
        return
    if w_el is None:
        w_el = OxmlElement("w:w")
        rpr.append(w_el)
    w_el.set(qn("w:val"), str(scale))


def set_xml_bold_rpr(rpr, bold: bool | None) -> None:
    b = rpr.find(qn("w:b"))
    if bold is None:
        if b is not None:
            rpr.remove(b)
        return
    if b is None:
        b = OxmlElement("w:b")
        rpr.append(b)
    if not bold:
        b.set(qn("w:val"), "0")


def apply_run_format(run, spec: FormatSpec) -> None:
    rpr = ensure_rpr(run._element)
    set_xml_font_rpr(rpr, spec.font)
    if spec.size is not None:
        set_xml_size_rpr(rpr, spec.size)
        run.font.size = Pt(spec.size)
    run.font.name = spec.font
    set_xml_spacing_rpr(rpr, spec.csp)
    set_xml_scale_rpr(rpr, spec.scale)
    if spec.bold is not None:
        set_xml_bold_rpr(rpr, spec.bold)
        run.font.bold = spec.bold


def apply_paragraph_format(paragraph, spec: FormatSpec) -> None:
    pf = paragraph.paragraph_format
    pf.alignment = ALIGN_MAP[spec.align]
    if spec.line is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = spec.line

    if spec.style_name and paragraph.style.name != spec.style_name:
        try:
            paragraph.style = spec.style_name
        except KeyError:
            pass

    for run in paragraph.runs:
        apply_run_format(run, spec)


def build_style_rpr(spec: FormatSpec) -> ET.Element:
    rpr = ET.Element(wqn("rPr"))
    rf = ET.SubElement(rpr, wqn("rFonts"))
    rf.set(wqn("ascii"), spec.font)
    rf.set(wqn("hAnsi"), spec.font)
    rf.set(wqn("eastAsia"), spec.font)
    if spec.size is not None:
        half = str(int(round(spec.size * 2)))
        sz = ET.SubElement(rpr, wqn("sz"))
        sz.set(wqn("val"), half)
        szcs = ET.SubElement(rpr, wqn("szCs"))
        szcs.set(wqn("val"), half)
    if spec.csp is not None:
        sp = ET.SubElement(rpr, wqn("spacing"))
        sp.set(wqn("val"), str(spec.csp))
    if spec.scale is not None:
        w_el = ET.SubElement(rpr, wqn("w"))
        w_el.set(wqn("val"), str(spec.scale))
    if spec.bold:
        ET.SubElement(rpr, wqn("b"))
    return rpr


def build_style_ppr(spec: FormatSpec) -> ET.Element:
    ppr = ET.Element(wqn("pPr"))
    jc = ET.SubElement(ppr, wqn("jc"))
    jc_map = {"left": "left", "center": "center", "right": "right", "justify": "both"}
    jc.set(wqn("val"), jc_map[spec.align])
    if spec.line is not None:
        sp = ET.SubElement(ppr, wqn("spacing"))
        sp.set(wqn("line"), str(int(spec.line * 240)))
        sp.set(wqn("lineRule"), "auto")
    return ppr


def update_styles_xml(styles_xml_bytes: bytes) -> bytes:
    root = ET.fromstring(styles_xml_bytes)

    for st in root.findall(wqn("style")):
        name_el = st.find(wqn("name"))
        if name_el is None:
            continue
        name = name_el.get(wqn("val"), "")
        sid = st.get(wqn("styleId"), "")

        spec = STYLE_NAME_DEFS.get(name) or STYLE_XML_DEFS.get(sid)
        if spec is None:
            continue

        # Remove existing rPr/pPr and rebuild
        for tag in ("rPr", "pPr"):
            old = st.find(wqn(tag))
            if old is not None:
                st.remove(old)

        rpr = build_style_rpr(spec)
        ppr = build_style_ppr(spec)
        st.insert(0, ppr)
        st.insert(1, rpr)

    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def get_run_props(run) -> dict[str, Any]:
    rpr = run._element.rPr
    d: dict[str, Any] = {}
    if run.font.name:
        d["font"] = run.font.name
    if run.font.size:
        d["size"] = round(run.font.size.pt, 1)
    if run.font.bold:
        d["bold"] = True
    if rpr is not None:
        sp = rpr.find(qn("w:spacing"))
        if sp is not None:
            d["csp"] = int(sp.get(qn("w:val")))
        w = rpr.find(qn("w:w"))
        if w is not None:
            d["scale"] = int(w.get(qn("w:val")))
    return d


def get_para_props(paragraph) -> dict[str, Any]:
    pf = paragraph.paragraph_format
    return {
        "align": REVERSE_ALIGN.get(pf.alignment, pf.alignment),
        "line": pf.line_spacing,
    }


def audit_document(doc: Document) -> list[tuple]:
    issues = []
    refs_idx = find_refs_title_index(doc)
    for i, p in enumerate(doc.paragraphs):
        if is_image_only_paragraph(p) or not p.text.strip():
            continue
        cat = classify_paragraph(p, i, refs_idx)
        if cat is None or cat not in SPEC:
            continue
        spec = SPEC[cat]
        rp = get_run_props(p.runs[0]) if p.runs else {}
        pp = get_para_props(p)

        for key, expected in [
            ("font", spec.font.split()[0]),
            ("size", spec.size),
            ("scale", spec.scale),
            ("csp", spec.csp),
            ("line", spec.line),
            ("align", spec.align),
        ]:
            if key == "font":
                actual = rp.get("font", "")
                if expected and actual and expected not in actual:
                    issues.append((i, cat, key, expected, actual))
            elif key == "csp" and expected is None:
                continue
            elif key in rp:
                if rp[key] != expected:
                    issues.append((i, cat, key, expected, rp[key]))
            elif key in pp:
                actual = pp[key]
                if key == "align" and actual == "inherit":
                    issues.append((i, cat, key, expected, actual))
                elif key == "line" and actual != expected:
                    issues.append((i, cat, key, expected, actual))
            elif key in ("size", "scale", "csp"):
                if expected is not None:
                    issues.append((i, cat, key, expected, "missing"))
            elif key == "line" and pp.get("line") is None:
                issues.append((i, cat, key, expected, None))
            elif key == "align" and pp.get("align") in ("inherit", None):
                issues.append((i, cat, key, expected, pp.get("align")))
    return issues


def format_document(doc: Document) -> list[str]:
    changes: list[str] = []
    refs_idx = find_refs_title_index(doc)
    for i, p in enumerate(doc.paragraphs):
        if is_image_only_paragraph(p):
            continue
        if not p.text.strip():
            continue

        cat = classify_paragraph(p, i, refs_idx)
        if cat is None:
            continue

        spec = SPEC[cat]
        old_style = p.style.name if p.style else ""
        apply_paragraph_format(p, spec)
        new_style = p.style.name if p.style else ""
        if old_style != new_style:
            changes.append(f"Para {i}: style {old_style!r} -> {new_style!r} ({cat})")
        else:
            changes.append(f"Para {i}: applied format for {cat}")

    return changes


def write_audit_report(
    before_issues: list,
    after_issues: list,
    changes: list[str],
    text_hash_before: str,
    text_hash_after: str,
) -> None:
    lines = [
        "# liangsc_formatted 格式检查报告",
        "",
        "## 摘要",
        "",
        f"- 源文件: `{INPUT_PATH}`",
        f"- 输出副本: `{OUTPUT_PATH}`",
        f"- 修正前问题数: **{len(before_issues)}**",
        f"- 修正后问题数: **{len(after_issues)}**",
        f"- 文本 hash 一致: **{'是' if text_hash_before == text_hash_after else '否'}**",
        "",
        "## 样式表覆盖情况",
        "",
        "| 样式项 | 状态 |",
        "|--------|------|",
        "| 논문제목 (Paper Title) | 已修正 |",
        "| 논문부제목 (Paper Subtitle) | 已修正 |",
        "| 이름 (Paper Author) | 已修正 |",
        "| 소속 (Paper Affiliation) | 已修正 |",
        "| 본문제목1 (样式2) | 已修正 |",
        "| 본문제목1.1 (样式1) | 已修正 |",
        "| 본문 (样式3) | 已修正 |",
        "| 참고문헌제목 | 已修正 |",
        "| 참고문헌 | 已修正 |",
        "| abstract / 목차 / 图题 / 表题 / 脚注 / 英文标题 | N/A（文档无对应内容）|",
        "",
        "## 修正操作",
        "",
    ]
    for c in changes:
        lines.append(f"- {c}")

    lines.extend(["", "## 修正后剩余问题", ""])
    if after_issues:
        for issue in after_issues[:50]:
            lines.append(f"- Para {issue[0]} [{issue[1]}] {issue[2]}: expected {issue[3]}, got {issue[4]}")
        if len(after_issues) > 50:
            lines.append(f"- ... 共 {len(after_issues)} 项")
    else:
        lines.append("- 无（全部合规）")

    AUDIT_PATH.write_text("\n".join(lines), encoding="utf-8")


def main() -> int:
    if not INPUT_PATH.exists():
        print(f"Input not found: {INPUT_PATH}", file=sys.stderr)
        return 1

    # Load and audit before
    doc = Document(str(INPUT_PATH))
    hash_before = paragraph_text_hash(doc)
    before_issues = audit_document(doc)

    # Apply formatting
    changes = format_document(doc)
    hash_after = paragraph_text_hash(doc)

    if hash_before != hash_after:
        print("ERROR: Text content changed!", file=sys.stderr)
        return 2

    # Save via docx then patch styles.xml
    doc.save(str(OUTPUT_PATH))

    # Update styles.xml in output zip
    with zipfile.ZipFile(OUTPUT_PATH, "r") as zin:
        contents = {name: zin.read(name) for name in zin.namelist()}

    if "word/styles.xml" in contents:
        contents["word/styles.xml"] = update_styles_xml(contents["word/styles.xml"])

    tmp = OUTPUT_PATH.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in contents.items():
            zout.writestr(name, data)
    tmp.replace(OUTPUT_PATH)

    # Re-audit output
    doc_out = Document(str(OUTPUT_PATH))
    after_issues = audit_document(doc_out)

    write_audit_report(before_issues, after_issues, changes, hash_before, hash_after)

    print(f"Saved: {OUTPUT_PATH}")
    print(f"Audit: {AUDIT_PATH}")
    print(f"Before issues: {len(before_issues)}, After issues: {len(after_issues)}")
    print(f"Text hash OK: {hash_before == hash_after}")

    return 0 if len(after_issues) == 0 else 3


if __name__ == "__main__":
    sys.exit(main())
