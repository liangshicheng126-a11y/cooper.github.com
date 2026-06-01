#!/usr/bin/env python3
"""Remove direct 장평/자간/행간/정렬 formatting added in the last batch pass."""

from __future__ import annotations

import hashlib
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn

DOC_PATH = Path(r"X:\A\liangsc_formatted_checked.docx")
WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def wqn(tag: str) -> str:
    return f"{{{WNS}}}{tag}"


def paragraph_text_hash(doc: Document) -> str:
    return hashlib.sha256("\n".join(p.text for p in doc.paragraphs).encode()).hexdigest()


def strip_ppr_spacing(p_pr: ET.Element) -> None:
    for tag in ("jc", "spacing", "ind"):
        el = p_pr.find(wqn(tag))
        if el is not None:
            p_pr.remove(el)


def strip_rpr_spacing(r_pr: ET.Element) -> None:
    for tag in ("spacing", "w", "kern", "fitText"):
        el = r_pr.find(wqn(tag))
        if el is not None:
            r_pr.remove(el)


def patch_document_xml(data: bytes) -> bytes:
    root = ET.fromstring(data)
    for p_pr in root.iter(wqn("pPr")):
        strip_ppr_spacing(p_pr)
    for r_pr in root.iter(wqn("rPr")):
        strip_rpr_spacing(r_pr)
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def patch_styles_xml(data: bytes) -> bytes:
    """Restore style-level defs: keep fonts/sizes/bold, drop line/scale/csp on 样式1/2/3."""
    root = ET.fromstring(data)
    strip_style_names = {"样式1", "样式2", "样式3", "Normal"}

    for st in root.findall(wqn("style")):
        name_el = st.find(wqn("name"))
        if name_el is None:
            continue
        name = name_el.get(wqn("val"), "")
        if name not in strip_style_names:
            continue
        p_pr = st.find(wqn("pPr"))
        if p_pr is not None:
            sp = p_pr.find(wqn("spacing"))
            if sp is not None:
                p_pr.remove(sp)
            if name in ("样式3", "Normal"):
                jc = p_pr.find(wqn("jc"))
                if jc is not None:
                    p_pr.remove(jc)
        r_pr = st.find(wqn("rPr"))
        if r_pr is not None:
            strip_rpr_spacing(r_pr)

    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def patch_footer_xml(data: bytes) -> bytes:
    root = ET.fromstring(data)
    for p_pr in root.iter(wqn("pPr")):
        jc = p_pr.find(wqn("jc"))
        if jc is not None:
            p_pr.remove(jc)
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def main() -> int:
    if not DOC_PATH.exists():
        print(f"Not found: {DOC_PATH}", file=sys.stderr)
        return 1

    hash_before = paragraph_text_hash(Document(str(DOC_PATH)))

    with zipfile.ZipFile(DOC_PATH, "r") as zin:
        contents = {name: zin.read(name) for name in zin.namelist()}

    contents["word/document.xml"] = patch_document_xml(contents["word/document.xml"])
    if "word/styles.xml" in contents:
        contents["word/styles.xml"] = patch_styles_xml(contents["word/styles.xml"])
    for key in list(contents):
        if key.startswith("word/footer") and key.endswith(".xml"):
            contents[key] = patch_footer_xml(contents[key])

    tmp = DOC_PATH.with_suffix(".revert.tmp.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in contents.items():
            zout.writestr(name, data)
    tmp.replace(DOC_PATH)

    hash_after = paragraph_text_hash(Document(str(DOC_PATH)))
    if hash_before != hash_after:
        print("ERROR: text changed", file=sys.stderr)
        return 2

    print(f"Reverted spacing/alignment direct format: {DOC_PATH}")
    print("Text hash OK")
    return 0


if __name__ == "__main__":
    sys.exit(main())
