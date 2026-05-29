# -*- coding: utf-8 -*-
from collections import Counter
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

path = r"X:\A\liangsc_formatted.docx"
out_path = r"X:\A\1\inspect_output.txt"
doc = Document(path)
lines = []


def log(msg=""):
    lines.append(msg)


def run_font_info(r):
    rpr = r._element.rPr
    east = None
    spacing = None
    if rpr is not None:
        if rpr.rFonts is not None:
            east = rpr.rFonts.get(qn("w:eastAsia"))
        sp = rpr.find(qn("w:spacing"))
        if sp is not None:
            spacing = sp.get(qn("w:val"))
    return r.font.name, east, r.font.size, r.font.bold, spacing


log("=== STYLES (paragraph) ===")
for s in doc.styles:
    if s.type == WD_STYLE_TYPE.PARAGRAPH:
        try:
            f = s.font
            log(f"{s.name}: font={f.name}, size={f.size}, bold={f.bold}")
        except Exception:
            pass

log("\n=== SECTIONS ===")
for i, sec in enumerate(doc.sections):
    log(
        f"Section {i}: page={sec.page_width.inches:.2f}x{sec.page_height.inches:.2f} in, "
        f"margins L={sec.left_margin.inches:.2f} R={sec.right_margin.inches:.2f} "
        f"T={sec.top_margin.inches:.2f} B={sec.bottom_margin.inches:.2f}"
    )

log("\n=== PARAGRAPHS (first 20 with text) ===")
count = 0
for i, p in enumerate(doc.paragraphs):
    if not p.text.strip() and count > 5:
        continue
    pf = p.paragraph_format
    log(
        f"P{i}: style={p.style.name} align={p.alignment} "
        f"line={pf.line_spacing} space_b={pf.space_before} space_a={pf.space_after} "
        f"first_indent={pf.first_line_indent}"
    )
    for j, r in enumerate(p.runs[:2]):
        fn, east, size, bold, spacing = run_font_info(r)
        log(
            f"  run{j}: font={fn} east={east} size={size} "
            f"bold={bold} char_spacing={spacing} text={r.text[:40]!r}"
        )
    if p.text.strip():
        log(f"  text: {p.text[:100]}")
        count += 1
    if count >= 20:
        break

log(f"\nTotal paragraphs: {len(doc.paragraphs)}")

styles_used = Counter(p.style.name for p in doc.paragraphs if p.text.strip())
log("\n=== STYLE USAGE ===")
for name, cnt in styles_used.most_common():
    log(f"  {name}: {cnt}")

log("\n=== CUSTOM PAPER STYLES ===")
for name in [
    "Paper Title",
    "Paper Subtitle",
    "Paper Author",
    "Paper Affiliation",
    "Paper H1",
    "Paper H2",
    "Paper H3",
    "Paper Ref",
    "Normal",
]:
    try:
        s = doc.styles[name]
        f = s.font
        pf = s.paragraph_format
        east = None
        spacing = None
        if s._element.rPr is not None:
            if s._element.rPr.rFonts is not None:
                east = s._element.rPr.rFonts.get(qn("w:eastAsia"))
            sp = s._element.rPr.find(qn("w:spacing"))
            if sp is not None:
                spacing = sp.get(qn("w:val"))
        log(
            f"{name}: font={f.name} east={east} size={f.size} bold={f.bold} "
            f"align={pf.alignment} line={pf.line_spacing} "
            f"space_b={pf.space_before} space_a={pf.space_after} "
            f"first_indent={pf.first_line_indent} char_spacing={spacing}"
        )
    except KeyError:
        log(f"{name}: (not found)")

with open(out_path, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))
print("Wrote", out_path)
