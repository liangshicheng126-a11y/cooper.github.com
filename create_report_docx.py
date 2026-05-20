# -*- coding: utf-8 -*-
"""Generate 111.docx bibliometric report for X:\\A."""
import json
import os
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT_PATH = r"X:\A\111.docx"
REPORT_JSON = r"X:\A\1\vos_data\output\analysis_report.json"
IMG_DIR = r"X:\A\1\vos_data\output"
SOURCE_DOC_HINT = "研究关注在文化消费背景下（X:\\A 文件夹 Word 文档）"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main():
    with open(REPORT_JSON, encoding="utf-8") as f:
        report = json.load(f)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("文化消费与视觉传达文献计量分析报告\n（VOSviewer 关键词共现与作者合作网络）")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph(f"生成日期：{date.today().isoformat()}")
    doc.add_paragraph(f"数据来源：{SOURCE_DOC_HINT}")

    add_heading(doc, "一、研究背景与文献来源", 1)
    doc.add_paragraph(
        "源文档围绕文化消费背景下非遗品牌视觉符号、文化认知—文化认同—购买偏好链条，"
        "以及 CNN-STING 混合模型在现代艺术设计特征识别中的应用。"
        "文档正文及嵌入的 Scopus 截图共识别出 2 篇核心文献，并以其参考文献扩展为 82 条记录的分析语料。"
    )

    add_heading(doc, "二、核心文献", 2)
    for i, p in enumerate(report["seed_papers"], 1):
        doc.add_paragraph(f"{i}. {p['title']}")
        doc.add_paragraph(f"   DOI: {p['doi']}")
        doc.add_paragraph(f"   作者: {', '.join(p['authors'])}")

    add_heading(doc, "三、软件环境", 1)
    add_bullets(
        doc,
        [
            "已安装 Java：Eclipse Temurin JRE 17（VOSviewer 运行依赖）",
            "已下载 VOSviewer 1.6.20 至 X:\\A\\1\\VOSviewer（便携版，解压后运行 VOSviewer.exe）",
            "导入数据：X:\\A\\1\\vos_data\\corpus.ris（RIS 格式，含题名、作者、关键词、机构、国家）",
            "分析类型：关键词共现（Co-occurrence）、作者/机构/国家合作（Co-authorship）",
        ],
    )

    add_heading(doc, "四、语料概况", 1)
    doc.add_paragraph(
        f"语料规模：共 {report['corpus_size']} 篇文献（2 篇种子文献 + 去重后的参考文献）。"
        "元数据通过 OpenAlex API 获取并转换为 VOSviewer 可读的 RIS 文件。"
    )

    add_heading(doc, "五、关键词共现分析", 1)
    doc.add_paragraph("在同一篇文献内同时出现的关键词构成共现关系；下表为高频关键词及典型共现对。")
    doc.add_paragraph("表1 高频关键词（文档频次）")
    t1 = doc.add_table(rows=1, cols=2)
    t1.style = "Table Grid"
    t1.rows[0].cells[0].text = "关键词"
    t1.rows[0].cells[1].text = "出现文献数"
    for kw, cnt in report["top_keywords"][:15]:
        row = t1.add_row().cells
        row[0].text = str(kw)
        row[1].text = str(cnt)

    doc.add_paragraph("表2 高频关键词共现对")
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "关键词对"
    t2.rows[0].cells[1].text = "共现次数"
    for pair, cnt in report["top_keyword_pairs"][:12]:
        row = t2.add_row().cells
        row[0].text = str(pair)
        row[1].text = str(cnt)

    kw_graph = next(g for g in report["graphs"] if g["name"] == "keywords")
    doc.add_paragraph(
        f"关键词网络：{kw_graph['nodes']} 个节点、{kw_graph['edges']} 条边。"
        "与源研究主题相关的节点包括 cultural heritage、brand management、structural equation modeling、"
        "intangible cultural heritage、visual communication、convolutional neural network 等，"
        "体现文化消费/非遗品牌与视觉智能设计两条研究线索的交叉。"
    )
    kw_img = report["images"]["keyword"]
    if os.path.isfile(kw_img):
        doc.add_paragraph("图1 关键词共现网络（VOSviewer 方法等效可视化）")
        doc.add_picture(kw_img, width=Inches(6.2))

    add_heading(doc, "六、作者合作网络分析", 1)
    au_graph = next(g for g in report["graphs"] if g["name"] == "authors")
    doc.add_paragraph(
        f"作者合作网络：{au_graph['nodes']} 位作者、{au_graph['edges']} 条合作关系。"
        "种子文献作者群：Zhao/Xie/Huang（韩国汉阳大学同一机构团队）与 Zhang Haiyan（兰州交通大学）"
        "在语料中形成相对独立的合作子群；扩展文献中营销、品牌、文化研究领域的合作团簇最为密集。"
    )
    doc.add_paragraph("表3 合作频次较高的作者（度中心性前12）")
    t3 = doc.add_table(rows=1, cols=2)
    t3.style = "Table Grid"
    t3.rows[0].cells[0].text = "作者"
    t3.rows[0].cells[1].text = "合作连接数"
    for name, deg in au_graph["top_nodes"]:
        row = t3.add_row().cells
        row[0].text = str(name)
        row[1].text = str(deg)

    au_img = report["images"]["author"]
    if os.path.isfile(au_img):
        doc.add_paragraph("图2 作者合作网络")
        doc.add_picture(au_img, width=Inches(6.2))

    add_heading(doc, "七、机构与国家合作网络", 1)
    for key, label, fig_title, img_key in [
        ("institutions", "机构", "图3 机构合作网络", "institution"),
        ("countries", "国家", "图4 国家合作网络", "country"),
    ]:
        g = next(x for x in report["graphs"] if x["name"] == key)
        doc.add_paragraph(
            f"{label}合作网络：{g['nodes']} 个节点、{g['edges']} 条边。"
        )
        top = "；".join(f"{n}({d})" for n, d in g["top_nodes"][:8])
        doc.add_paragraph(f"主要节点：{top}")
        img = report["images"][img_key]
        if os.path.isfile(img):
            doc.add_paragraph(fig_title)
            doc.add_picture(img, width=Inches(6.0))

    add_heading(doc, "八、主要结论", 1)
    conclusions = [
        "文献主题呈现“文化消费与非遗品牌视觉符号”与“现代艺术视觉传达智能模型（CNN-STING）”两条主线，"
        "在关键词层面通过 brand management、cultural heritage、visual communication、AI 等节点产生关联。",
        "关键词共现核心簇为 advertising–business–marketing–psychology，说明扩展参考文献以品牌与消费者行为研究为主。",
        "作者合作呈现典型的小团队结构：种子论文团队内部合作紧密，跨团队国际合作在扩展语料中由多国机构边体现。",
        "建议在 VOSviewer 中进一步设置最小共现频次、密度可视化与聚类着色，以区分文化研究与计算设计子领域。",
    ]
    add_bullets(doc, conclusions)

    add_heading(doc, "九、复现步骤（VOSviewer）", 1)
    steps = [
        "启动 X:\\A\\1\\VOSviewer 中的 VOSviewer.exe（需已安装 Java 17）。",
        "Create → Create a map based on bibliographic data → Read data from bibliographic database files → RIS.",
        "选择 X:\\A\\1\\vos_data\\corpus.ris，分析类型选 Co-occurrence（Author keywords）或 Co-authorship（Authors/Organizations/Countries）。",
        "设置最小出现次数（如 keywords ≥2），运行 Continue → Finish，即可得到与本文附图一致的交互式网络图。",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}")

    doc.add_paragraph(
        "\n附件数据目录：X:\\A\\1\\vos_data\\（corpus.ris、corpus.json、output/*.png）"
    )

    doc.save(OUT_PATH)
    print("Saved:", OUT_PATH)


if __name__ == "__main__":
    main()
