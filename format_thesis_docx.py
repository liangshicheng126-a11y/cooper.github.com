# -*- coding: utf-8 -*-
"""Format liangsc.docx per teacher Oral Presentation style sheet."""
from __future__ import annotations

import re
import shutil
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

INPUT = Path(r"X:\A\liangsc.docx")
OUTPUT = Path(r"X:\A\liangsc_formatted.docx")

FONT_GOTHIC = "Nanum Gothic"
FONT_GOTHIC_EB = "Nanum Gothic ExtraBold"
FONT_GOTHIC_B = "Nanum Gothic Bold"
FONT_MYUNGJO = "Nanum Myeongjo"

REQUIRED_FONTS = [FONT_GOTHIC, FONT_GOTHIC_EB, FONT_MYUNGJO]

MARGINS = dict(top=Cm(2.5), bottom=Cm(2.5), left=Cm(2.5), right=Cm(2.5))
COL_SPACE_TWIPS = "425"  # ~0.75 cm between columns

RE_H1 = re.compile(r"^\d+\.\s*\S")
RE_H2 = re.compile(r"^\d+\.\d+\.?\s*\S")
RE_H3 = re.compile(r"^\d+\.\d+\.\d+\.?\s*\S")
RE_REF_HEAD = re.compile(r"^(참고문헌|References)\s*$", re.I)


@dataclass(frozen=True)
class StyleSpec:
    font: str
    size_pt: float
    scale_pct: int  # 장평 horizontal scale
    char_spacing: int  # twips (1/20 pt); 0 = standard
    line_pct: float  # e.g. 1.6 = 160%
    align: WD_ALIGN_PARAGRAPH
    bold: bool = False


STYLE_SPEC: dict[str, StyleSpec] = {
    "title": StyleSpec(FONT_GOTHIC_EB, 18, 90, -10, 1.3, WD_ALIGN_PARAGRAPH.CENTER, True),
    "subtitle": StyleSpec(FONT_GOTHIC, 11, 100, -10, 1.3, WD_ALIGN_PARAGRAPH.CENTER),
    "author": StyleSpec(FONT_GOTHIC, 12, 95, -5, 1.6, WD_ALIGN_PARAGRAPH.CENTER),
    "affiliation": StyleSpec(FONT_GOTHIC, 9.5, 95, -5, 1.6, WD_ALIGN_PARAGRAPH.CENTER),
    "h1": StyleSpec(FONT_MYUNGJO, 13, 100, -5, 1.6, WD_ALIGN_PARAGRAPH.LEFT),
    "h2": StyleSpec(FONT_GOTHIC, 11, 100, -5, 1.6, WD_ALIGN_PARAGRAPH.LEFT),
    "h3": StyleSpec(FONT_GOTHIC_B, 10, 100, -5, 1.6, WD_ALIGN_PARAGRAPH.LEFT, True),
    "body": StyleSpec(FONT_MYUNGJO, 10, 100, -5, 1.6, WD_ALIGN_PARAGRAPH.JUSTIFY),
    "ref_heading": StyleSpec(FONT_GOTHIC_EB, 11, 100, -5, 1.6, WD_ALIGN_PARAGRAPH.LEFT, True),
    "ref_entry": StyleSpec(FONT_MYUNGJO, 9, 100, 0, 1.6, WD_ALIGN_PARAGRAPH.LEFT),
    "figure_caption": StyleSpec(FONT_GOTHIC, 8, 95, -5, 1.3, WD_ALIGN_PARAGRAPH.CENTER),
    "table_caption": StyleSpec(FONT_GOTHIC, 8, 95, -5, 1.3, WD_ALIGN_PARAGRAPH.LEFT),
}


def check_fonts() -> list[str]:
    missing = []
    try:
        import winreg

        key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts")
        installed = " ".join(winreg.EnumValue(key, i)[0] for i in range(winreg.QueryInfoKey(key)[1])).lower()
        for font in REQUIRED_FONTS:
            if font.lower().replace(" ", "") not in installed.replace(" ", ""):
                missing.append(font)
    except OSError:
        pass
    return missing


def set_char_spacing(run, twips: int) -> None:
    rpr = run._element.get_or_add_rPr()
    spacing = rpr.find(qn("w:spacing"))
    if twips == 0:
        if spacing is not None:
            rpr.remove(spacing)
        return
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        rpr.append(spacing)
    spacing.set(qn("w:val"), str(twips))


def set_text_scale(run, pct: int) -> None:
    rpr = run._element.get_or_add_rPr()
    ts = rpr.find(qn("w:textScale"))
    if pct == 100:
        if ts is not None:
            rpr.remove(ts)
        return
    if ts is None:
        ts = OxmlElement("w:textScale")
        rpr.append(ts)
    ts.set(qn("w:val"), str(pct))


def set_run_style(run, spec: StyleSpec) -> None:
    run.font.name = spec.font
    run.font.size = Pt(spec.size_pt)
    run.font.bold = spec.bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for tag in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{tag}"), spec.font)
    set_char_spacing(run, spec.char_spacing)
    set_text_scale(run, spec.scale_pct)


def apply_paragraph_format(p: Paragraph, spec: StyleSpec) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = spec.line_pct
    pf.alignment = spec.align
    pf.first_line_indent = Pt(0)


def replace_paragraph_text(p: Paragraph, new_text: str) -> None:
    if not p.runs:
        p.add_run(new_text)
        return
    p.runs[0].text = new_text
    for run in p.runs[1:]:
        run.text = ""


def normalize_heading_text(text: str, kind: str) -> str:
    text = text.strip()
    if kind == "h1":
        m = re.match(r"^(\d+)\.\s*(.+)$", text)
        if m:
            return f"{m.group(1)}. {m.group(2).strip()}"
    if kind == "h2":
        m = re.match(r"^(\d+\.\d+)\.?\s*(.+)$", text)
        if m:
            return f"{m.group(1)}. {m.group(2).strip()}"
    if kind == "h3":
        m = re.match(r"^(\d+\.\d+\.\d+)\.?\s*(.+)$", text)
        if m:
            return f"{m.group(1)}. {m.group(2).strip()}"
    return text


def classify_paragraph(
    style_name: str,
    text: str,
    bold_runs: bool,
    in_refs: bool,
) -> str:
    text = text.strip()
    if not text:
        return "empty"
    if in_refs:
        return "ref_entry"
    if style_name == "Paper Title":
        return "title"
    if style_name == "Paper Subtitle":
        return "subtitle"
    if style_name == "Paper Author":
        return "author"
    if style_name == "Paper Affiliation":
        return "affiliation"
    if RE_REF_HEAD.match(text):
        return "ref_heading"
    if RE_H3.match(text) or style_name == "Paper H3":
        return "h3"
    if RE_H2.match(text) or style_name in {"样式1", "Paper H2"}:
        return "h2"
    if RE_H1.match(text) or style_name in {"样式2", "Paper H1"}:
        return "h1"
    if bold_runs and RE_H2.match(text):
        return "h2"
    if text.startswith("<표") or text.startswith("[그림"):
        return "table_caption" if text.startswith("<") else "figure_caption"
    return "body"


def format_paragraph(p: Paragraph, kind: str) -> None:
    if kind == "empty":
        return
    spec = STYLE_SPEC.get(kind)
    if spec is None:
        spec = STYLE_SPEC["body"]
        kind = "body"

    text = p.text.strip()
    if kind in {"h1", "h2", "h3"}:
        text = normalize_heading_text(text, kind)
        replace_paragraph_text(p, text)

    apply_paragraph_format(p, spec)
    if not p.runs:
        run = p.add_run(text)
        set_run_style(run, spec)
        return
    for run in p.runs:
        set_run_style(run, spec)


def insert_section_break_after(paragraph: Paragraph) -> None:
    ppr = paragraph._element.get_or_add_pPr()
    if ppr.find(qn("w:sectPr")) is not None:
        return
    sect_pr = OxmlElement("w:sectPr")
    ppr.append(sect_pr)


def set_section_columns(section, num: int = 2) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    if num == 1:
        if cols.get(qn("w:num")):
            cols.attrib.pop(qn("w:num"), None)
        return
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), COL_SPACE_TWIPS)


def clear_footer(section) -> None:
    section.footer.is_linked_to_previous = False
    ftr = section.footer._element
    for child in list(ftr):
        if child.tag == qn("w:p"):
            ftr.remove(child)


def add_page_number_footer(section) -> None:
    section.footer.is_linked_to_previous = False
    ftr = section.footer._element
    for child in list(ftr):
        if child.tag == qn("w:p"):
            ftr.remove(child)
    p = section.footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spec = STYLE_SPEC["body"]

    r1 = p.add_run("- ")
    set_run_style(r1, StyleSpec(FONT_MYUNGJO, 10, 100, -5, 1.6, WD_ALIGN_PARAGRAPH.CENTER))

    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    fr = OxmlElement("w:r")
    ft = OxmlElement("w:t")
    ft.text = "1"
    fr.append(ft)
    fld.append(fr)
    p._element.append(fld)

    r2 = p.add_run(" -")
    set_run_style(r2, StyleSpec(FONT_MYUNGJO, 10, 100, -5, 1.6, WD_ALIGN_PARAGRAPH.CENTER))


def find_affiliation_index(doc: Document) -> Optional[int]:
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Paper Affiliation":
            return i
    for i, p in enumerate(doc.paragraphs):
        if "대구가톨릭" in p.text or "대학원" in p.text:
            return i
    return None


def setup_layout(doc: Document) -> None:
    for sec in doc.sections:
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = MARGINS["top"]
        sec.bottom_margin = MARGINS["bottom"]
        sec.left_margin = MARGINS["left"]
        sec.right_margin = MARGINS["right"]

    aff_idx = find_affiliation_index(doc)
    if aff_idx is not None:
        insert_section_break_after(doc.paragraphs[aff_idx])

    sections = doc.sections
    if len(sections) >= 1:
        set_section_columns(sections[0], 1)
        clear_footer(sections[0])
    if len(sections) >= 2:
        set_section_columns(sections[1], 2)
        clear_footer(sections[1])
        add_page_number_footer(sections[1])
    elif len(sections) == 1:
        set_section_columns(sections[0], 2)
        clear_footer(sections[0])
        add_page_number_footer(sections[0])


def update_normal_style(doc: Document) -> None:
    spec = STYLE_SPEC["body"]
    normal = doc.styles["Normal"]
    normal.font.name = spec.font
    normal.font.size = Pt(spec.size_pt)
    rpr = normal._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for tag in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{tag}"), spec.font)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = spec.line_pct


def process_document(doc: Document) -> None:
    in_refs = False
    kinds: list[str] = []

    for p in doc.paragraphs:
        bold = any(r.bold for r in p.runs if r.text.strip())
        kind = classify_paragraph(p.style.name, p.text, bold, in_refs)
        if kind == "ref_heading":
            in_refs = True
        kinds.append(kind)
        format_paragraph(p, kind)

    return kinds


def main() -> None:
    if not INPUT.exists():
        raise FileNotFoundError(INPUT)

    missing = check_fonts()
    if missing:
        print("Warning: fonts may be missing:", ", ".join(missing))
        print("Download: https://hangeul.naver.com/font")

    backup = INPUT.with_name(
        INPUT.stem + "_backup_" + datetime.now().strftime("%Y%m%d_%H%M%S") + INPUT.suffix
    )
    shutil.copy2(INPUT, backup)

    doc = Document(str(INPUT))
    update_normal_style(doc)
    kinds = process_document(doc)
    setup_layout(doc)

    doc.save(str(OUTPUT))
    print(f"Backup: {backup}")
    print(f"Saved:  {OUTPUT}")
    print(f"Sections: {len(doc.sections)}")
    print(f"Paragraph kinds: { {k: kinds.count(k) for k in set(kinds) if k != 'empty'} }")


if __name__ == "__main__":
    main()
